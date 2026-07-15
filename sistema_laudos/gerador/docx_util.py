"""
Utilitários de baixo nível para o documento Word (.docx).

Concentra o "encanamento" que satisfaz os critérios de aceite da Fase 1:
  * Sumário automático (campo TOC que o Word popula ao abrir);
  * Numeração de página contínua (campo PAGE no rodapé) e atualização
    automática de campos na abertura;
  * Compressão de imagem no pipeline (Pillow, máx. 1600px, JPEG ~80%);
  * Legenda de figura numerada automaticamente;
  * Fontes e cores da marca (design tokens).
"""
from __future__ import annotations

import json
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps

# --------------------------------------------------------------------------
# Design tokens
# --------------------------------------------------------------------------
_TOKENS_PATH = Path(__file__).resolve().parent.parent / "design" / "tokens.json"


def carregar_tokens() -> dict:
    return json.loads(_TOKENS_PATH.read_text(encoding="utf-8"))


def hexnum(cor: str) -> str:
    """'#171f3d' -> '171F3D' (formato exigido pelo OOXML)."""
    return cor.lstrip("#").upper()


# --------------------------------------------------------------------------
# Campos do Word (TOC, PAGE, updateFields)
# --------------------------------------------------------------------------
def _campo(paragraph, instrucao: str, placeholder: str = ""):
    """Insere um campo do Word ( { instrucao } ) em um parágrafo."""
    r1 = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    r1._r.append(fld_begin)

    r2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instrucao
    r2._r.append(instr)

    r3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(fld_sep)

    if placeholder:
        paragraph.add_run(placeholder)

    r5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5._r.append(fld_end)


def inserir_sumario(paragraph):
    """Campo TOC (níveis 1–3). O Word popula ao abrir / com F9."""
    _campo(
        paragraph,
        ' TOC \\o "1-3" \\h \\z \\u ',
        placeholder="Sumário — atualize os campos (Ctrl+A, F9) ao abrir.",
    )


def atualizar_campos_ao_abrir(document):
    """Faz o Word atualizar TODOS os campos (sumário, paginação) ao abrir."""
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


# --------------------------------------------------------------------------
# Cabeçalho e rodapé (numeração de página contínua)
# --------------------------------------------------------------------------
def configurar_cabecalho_rodape(document, tokens, numero_laudo: str):
    """Cabeçalho fixo (CREA/IBAPE) e rodapé com numeração contínua.
    A primeira página (capa) fica limpa via 'different first page'."""
    section = document.sections[0]
    section.different_first_page_header_footer = True
    aplicar_cabecalho_rodape(section, tokens, numero_laudo)


def secao_sem_cabecalho(section):
    """Desliga cabeçalho/rodapé herdados numa seção de páginas-imagem
    (o template full-bleed carrega o próprio cabeçalho/rodapé)."""
    for parte in (section.header, section.footer):
        parte.is_linked_to_previous = False
        for p in list(parte.paragraphs):
            p.clear()


def aplicar_cabecalho_rodape(section, tokens, numero_laudo: str):
    """Aplica o cabeçalho/rodapé de miolo a uma seção (desvinculada)."""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    fonte_corpo = tokens["tipografia"]["corpo"]
    cinza = hexnum(tokens["marca"]["cinza"])
    terracota = hexnum(tokens["marca"]["terracota"])

    # ---- Cabeçalho (páginas de miolo) ----
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(f"Laudo de Engenharia Nº {numero_laudo}    ·    "
                     "CREA/MG 173201    ·    IBAPE/MG 1221")
    run.font.name = fonte_corpo
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(cinza)
    _linha_inferior(hp, cinza)

    # ---- Rodapé (numeração contínua "Página X") ----
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Página ")
    r.font.name = fonte_corpo
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(terracota)
    _campo(fp, " PAGE ")
    r2 = fp.add_run(" de ")
    r2.font.name = fonte_corpo
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor.from_string(terracota)
    _campo(fp, " NUMPAGES ")
    # aplica cor/fonte aos runs de campo do rodapé
    for run in fp.runs:
        run.font.name = fonte_corpo
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(terracota)


def _linha_inferior(paragraph, cor_hex: str):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), cor_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


# --------------------------------------------------------------------------
# Tabelas — sombreamento de célula (matrizes de risco)
# --------------------------------------------------------------------------
def sombrear_celula(cell, cor_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)


# --------------------------------------------------------------------------
# Página-imagem full-bleed (capítulos renderizados dos templates do design)
# --------------------------------------------------------------------------
# ids de wp:docPr precisam ser únicos no documento (vários por laudo agora:
# capa + páginas de capítulo + caixas de nº de página). Usa o alocador do
# próprio python-docx (part.next_id varre todos os @id do documento) — o
# mesmo que add_picture usa para as figuras inline, então nunca colide.
def _novo_docpr_id(paragraph) -> int:
    return paragraph.part.next_id


def inserir_imagem_pagina_inteira(paragraph, caminho_img: str | Path,
                                  largura=None, altura=None):
    """Insere uma imagem ancorada à PÁGINA, cobrindo-a por inteiro (full-bleed).

    Usa `wp:anchor` posicionado em (0,0) relativo à página, atrás do texto —
    independe das margens da seção e não ocupa espaço no fluxo (não há risco
    de a linha "não caber" e empurrar uma página em branco). O parágrafo deve
    estar na página desejada (ex.: primeiro parágrafo da capa).
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Cm

    did = _novo_docpr_id(paragraph)
    largura = largura if largura is not None else Cm(21.0)   # A4
    altura = altura if altura is not None else Cm(29.7)
    rid, _ = paragraph.part.get_or_add_image(str(caminho_img))
    nome = Path(caminho_img).name
    xml = (
        f'<w:drawing {nsdecls("w", "wp", "a", "pic", "r")}>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        ' relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1"'
        ' allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{int(largura)}" cy="{int(altura)}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{did}" name="{nome}"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        f'<pic:nvPicPr><pic:cNvPr id="{did}" name="{nome}"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        '<pic:spPr><a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{int(largura)}" cy="{int(altura)}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    )
    paragraph.add_run()._r.append(parse_xml(xml))


def inserir_numero_pagina_flutuante(paragraph, x_emu: int, y_emu: int,
                                    fonte: str, tamanho_pt: float,
                                    cor_hex: str, sufixo: str = " | Pág.",
                                    largura_emu: int = 1080000,
                                    altura_emu: int = 240000):
    """Nº de página VIVO sobre uma página-imagem, no visual do design.

    Caixa de texto flutuante ancorada à página, alinhada à direita, cujo
    conteúdo é o campo PAGE + sufixo. (x_emu, y_emu) é o canto superior
    DIREITO do slot medido no render do template (data-dc-medir="pagina");
    a caixa se estende `largura_emu` para a esquerda.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    did = _novo_docpr_id(paragraph)
    # python-docx não registra o prefixo wps (wordprocessingShape): declarar à mão
    ns_wps = ('xmlns:wps='
              '"http://schemas.microsoft.com/office/word/2010/wordprocessingShape"')
    xml = (
        f'<w:drawing {nsdecls("w", "wp", "a", "r")} {ns_wps}>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        ' relativeHeight="10" behindDoc="0" locked="0" layoutInCell="1"'
        ' allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x_emu - largura_emu}'
        '</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}'
        '</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{largura_emu}" cy="{altura_emu}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{did}" name="pagina{did}"/>'
        '<a:graphic><a:graphicData'
        ' uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        f'<wps:cNvPr id="{did}" name="pagina{did}"/><wps:cNvSpPr txBox="1"/>'
        '<wps:spPr><a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{largura_emu}" cy="{altura_emu}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '<a:noFill/><a:ln><a:noFill/></a:ln></wps:spPr>'
        '<wps:txbx><w:txbxContent>'
        '<w:p><w:pPr><w:jc w:val="right"/><w:spacing w:after="0"/><w:rPr>'
        f'<w:rFonts w:ascii="{fonte}" w:hAnsi="{fonte}"/><w:b/>'
        f'<w:color w:val="{cor_hex}"/><w:sz w:val="{int(tamanho_pt * 2)}"/>'
        '</w:rPr></w:pPr>'
        f'<w:fldSimple w:instr=" PAGE "><w:r><w:rPr>'
        f'<w:rFonts w:ascii="{fonte}" w:hAnsi="{fonte}"/><w:b/>'
        f'<w:color w:val="{cor_hex}"/><w:sz w:val="{int(tamanho_pt * 2)}"/>'
        '</w:rPr><w:t>0</w:t></w:r></w:fldSimple>'
        '<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{fonte}" w:hAnsi="{fonte}"/><w:b/>'
        f'<w:color w:val="{cor_hex}"/><w:sz w:val="{int(tamanho_pt * 2)}"/>'
        f'</w:rPr><w:t xml:space="preserve">{sufixo}</w:t></w:r>'
        '</w:p>'
        '</w:txbxContent></wps:txbx>'
        '<wps:bodyPr rot="0" wrap="none" lIns="0" tIns="0" rIns="0" bIns="0"'
        ' anchor="t"><a:noAutofit/></wps:bodyPr>'
        '</wps:wsp>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    )
    paragraph.add_run()._r.append(parse_xml(xml))


PX_PARA_EMU = 360000 * 21.0 / 794  # px CSS (página 794 px = 21 cm) -> EMU


# --------------------------------------------------------------------------
# Imagens — compressão no pipeline + placeholder
# --------------------------------------------------------------------------
def comprimir_imagem(origem: str | Path, destino: str | Path,
                     lado_max: int = 1600, qualidade: int = 80) -> Path:
    """Redimensiona (máx. `lado_max` no maior lado) e recomprime JPEG."""
    origem, destino = Path(origem), Path(destino)
    destino.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(origem) as img:
        img = ImageOps.exif_transpose(img)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        img.thumbnail((lado_max, lado_max), Image.LANCZOS)
        img.save(destino, "JPEG", quality=qualidade, optimize=True)
    return destino


def gerar_placeholder(destino: str | Path, texto: str,
                      largura: int = 1200, altura: int = 800) -> Path:
    """Gera uma imagem-marcador (para fixture sem foto real)."""
    destino = Path(destino)
    destino.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (largura, altura), "#e4e5e5")
    d = ImageDraw.Draw(img)
    d.rectangle([8, 8, largura - 8, altura - 8], outline="#969a9a", width=3)
    try:
        fonte = ImageFont.truetype("DejaVuSans.ttf", 34)
        fonte_p = ImageFont.truetype("DejaVuSans.ttf", 24)
    except OSError:
        fonte = ImageFont.load_default()
        fonte_p = fonte
    d.text((largura / 2, altura / 2 - 30), "FOTO DE EXEMPLO",
           fill="#5e6161", anchor="mm", font=fonte)
    d.text((largura / 2, altura / 2 + 20), "(substituir pela foto da vistoria)",
           fill="#767a7a", anchor="mm", font=fonte_p)
    d.text((largura / 2, altura - 60), texto[:80],
           fill="#45494a", anchor="mm", font=fonte_p)
    img.save(destino, "JPEG", quality=80, optimize=True)
    return destino
