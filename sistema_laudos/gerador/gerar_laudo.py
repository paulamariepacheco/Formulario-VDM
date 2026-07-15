"""
Motor de geração do Laudo de Engenharia (.docx) — Fase 1.

Uso:
    python -m gerador.gerar_laudo dados/laudo_50_2026.json saida/Laudo-50-2026.docx

Monta um documento único com: capa, sumário automático, capítulos fixos,
diagnóstico por sistema construtivo, figuras numeradas/comprimidas, matrizes
IBAPE/GUT calculadas, plano de intervenção, conclusões e encerramento —
com numeração de página contínua e cabeçalho CREA/IBAPE.
"""
from __future__ import annotations

import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from . import render_paginas, risco, textos
from .docx_util import (
    PX_PARA_EMU,
    aplicar_cabecalho_rodape,
    atualizar_campos_ao_abrir,
    carregar_tokens,
    comprimir_imagem,
    configurar_cabecalho_rodape,
    gerar_placeholder,
    hexnum,
    inserir_imagem_pagina_inteira,
    inserir_numero_pagina_flutuante,
    inserir_sumario,
    secao_sem_cabecalho,
    sombrear_celula,
)
from .fonte_dados import carregar_de_json
from .modelos import Laudo

TIPO_ROTULO = textos.TIPO_ROTULO  # rótulos de exibição (movidos p/ textos.py)


class GeradorLaudo:
    def __init__(self, laudo: Laudo, saida: Path):
        self.laudo = laudo
        self.saida = Path(saida)
        self.tokens = carregar_tokens()
        self.doc = Document()
        self.dir_img = self.saida.parent / "img"
        self._fig_seq = 0
        # páginas-imagem pré-renderizadas: chave -> (png, rects medidos)
        self._paginas: dict[str, tuple[Path, dict]] = {}

    @staticmethod
    def _sentenca(texto: str) -> str:
        """Capitaliza apenas a inicial ('impermeabilização e drenagem' -> ...)."""
        return texto[:1].upper() + texto[1:] if texto else texto

    def _estilos(self):
        t = self.tokens
        navy = hexnum(t["marca"]["azul"])
        chumbo = hexnum(t["marca"]["chumbo"])
        corpo = hexnum(t["texto"]["corpo"])
        f_titulo = t["tipografia"]["titulo"]
        f_corpo = t["tipografia"]["corpo"]

        normal = self.doc.styles["Normal"]
        normal.font.name = f_corpo
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor.from_string(corpo)
        pf = normal.paragraph_format
        pf.line_spacing = 1.4
        pf.space_after = Pt(8)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for nome, tam, cor in (("Heading 1", 17, navy),
                               ("Heading 2", 14, navy),
                               ("Heading 3", 12, chumbo),
                               ("Title", 26, navy)):
            st = self.doc.styles[nome]
            st.font.name = f_titulo
            st.font.size = Pt(tam)
            st.font.bold = nome != "Title"
            st.font.color.rgb = RGBColor.from_string(cor)

        for section in self.doc.sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.2)

    # -------------------------------------------------------------- helpers
    def _p(self, texto="", *, align=None, size=None, bold=False,
           italic=False, cor=None, space_after=None):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if texto:
            run = p.add_run(texto)
            run.bold = bold
            run.italic = italic
            if size:
                run.font.size = Pt(size)
            if cor:
                run.font.color.rgb = RGBColor.from_string(hexnum(cor))
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    def _campo_estruturado(self, rotulo: str, valor: str):
        if not valor:
            return
        p = self.doc.add_paragraph()
        r = p.add_run(f"{rotulo}: ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(hexnum(self.tokens["marca"]["azul"]))
        p.add_run(valor)

    def _quebra_pagina(self):
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------- páginas-imagem (templates)
    def _preparar_paginas(self):
        """Renderiza TODAS as páginas-template do laudo numa única sessão de
        Chromium (capa + capítulos 2, 3, 4, 7, 9, 10) antes da montagem.

        A numeração de figuras é determinística e calculada aqui: mapa (01) e
        fachada (02) na página de Vistoria, fotos avulsas e de anomalias na
        sequência, e as três figuras do Referencial por último. Se Playwright/
        Chromium não estiverem disponíveis, `self._paginas` fica vazio e a
        montagem degrada para os capítulos nativos (texto Word), com aviso.
        """
        lau = self.laudo
        rp = render_paginas
        try:
            with rp.SessaoRender() as sr:
                self.dir_img.mkdir(parents=True, exist_ok=True)
                png, _ = sr.render(rp.montar_html_capa(lau, self.tokens),
                                   self.dir_img / "capa.png")
                self._paginas["capa"] = (png, {})

                # numeração de figuras (ordem de aparição no documento)
                ids_anom = {an.id for an in lau.anomalias}
                n_avulsas = sum(1 for f in lau.fotos if not f.anomalia_id)
                n_de_anom = sum(1 for f in lau.fotos
                                if f.anomalia_id in ids_anom)
                self._fig_seq = 2          # 01=mapa, 02=fachada; avulsas em 03+
                base_ref = 2 + n_avulsas + n_de_anom

                mapa = Path(lau.mapa_localizacao or "")
                if not (lau.mapa_localizacao and mapa.exists()):
                    mapa = gerar_placeholder(
                        self.dir_img / "mapa-localizacao.jpg",
                        "Mapa de localização — captura do Google Maps",
                        1364, 556)
                fachada = Path(lau.foto_fachada or "")
                if not (lau.foto_fachada and fachada.exists()):
                    fachada = gerar_placeholder(
                        self.dir_img / "foto-fachada.jpg",
                        "Foto da fachada", 900, 696)

                campos_v, raw_v = rp.campos_vistoria(
                    lau, mapa.resolve().as_uri(), fachada.resolve().as_uri(),
                    fig_mapa=1, fig_fachada=2)
                campos_m, raw_m = rp.campos_metodos()
                campos_e, variantes_e = rp.campos_encerramento(lau)
                paginas = [
                    ("motivacao", 2, rp.campos_motivacao(lau), None, None),
                    ("vistoria", 3, campos_v, raw_v, None),
                    ("metodos", 4, campos_m, raw_m, None),
                    ("especificacoes", 7, None, None, None),
                    ("referencial", 9,
                     rp.campos_referencial(base_ref + 1, base_ref + 2,
                                           base_ref + 3), None, None),
                    ("encerramento", 10, campos_e, None, variantes_e),
                ]
                for chave, num, campos, raw, variantes in paginas:
                    html = rp.montar_html_capitulo(
                        chave, lau, self.tokens, num,
                        campos_extra=campos, raw=raw, variantes=variantes)
                    png, rects = sr.render(
                        html, self.dir_img / f"pag-{chave}.png",
                        medir=["pagina"])
                    self._paginas[chave] = (png, rects)
        except rp.RenderIndisponivel as exc:
            print(f"[aviso] capítulos nativos (render indisponível: {exc})")
            self._paginas = {}
            self._fig_seq = 0

    def _secao_imagem(self):
        """Nova seção (sempre) para UMA página-imagem: sem cabeçalho/rodapé
        do Word (o template full-bleed traz os próprios) e garante que cada
        página-template comece em página nova."""
        from docx.enum.section import WD_SECTION
        s = self.doc.add_section(WD_SECTION.NEW_PAGE)
        secao_sem_cabecalho(s)
        self._em_imagem = True
        return s

    def _garantir_miolo(self):
        """Volta (se preciso) para uma seção de miolo nativo, reaplicando o
        cabeçalho/rodapé padrão. Idempotente: em miolo, não faz nada."""
        if not getattr(self, "_em_imagem", False):
            return
        from docx.enum.section import WD_SECTION
        s = self.doc.add_section(WD_SECTION.NEW_PAGE)
        aplicar_cabecalho_rodape(s, self.tokens, self.laudo.numero)
        self._em_imagem = False
        return s

    def _pagina_imagem(self, chave: str, titulo_toc: str | None = None):
        """Insere uma página-template pré-renderizada na página corrente.

        `titulo_toc`: heading nível 1 invisível (1pt, branco) que ancora o
        capítulo no Sumário e mantém a numeração — o título visível é o do
        próprio template. O slot de nº de página medido no render recebe uma
        caixa flutuante com o campo PAGE (numeração viva, visual do design).
        """
        png, rects = self._paginas[chave]
        if titulo_toc:
            h = self.doc.add_heading(titulo_toc, level=1)
            for run in h.runs:
                run.font.size = Pt(1)
                run.font.color.rgb = RGBColor.from_string("FFFFFF")
            pf = h.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        inserir_imagem_pagina_inteira(p, png)
        if "pagina" in rects:
            x, y, w, _alt = rects["pagina"]
            inserir_numero_pagina_flutuante(
                p,
                x_emu=int((x + w) * PX_PARA_EMU),
                y_emu=int(y * PX_PARA_EMU),
                fonte="IBM Plex Mono",
                tamanho_pt=8.25,   # 11 px CSS a 96 dpi
                cor_hex=hexnum(self.tokens["marca"]["terracota"]))

    # ----------------------------------------------------------------- capa
    def _capa(self):
        """Capa como página-imagem renderizada do template do Claude Design.

        Fidelidade visual total: o template A4 (LaudoCapa) é resolvido com os
        dados do laudo e renderizado em PNG ~300 dpi, inserido full-bleed
        (âncora à página, sem margens). Sem foto de capa, o próprio template
        cai na variante `azul-marinho`. Se Playwright/Chromium não estiverem
        disponíveis, degrada para a capa nativa (texto Word) com aviso.
        """
        if "capa" not in self._paginas:
            self._capa_nativa()
            return
        png, _ = self._paginas["capa"]
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        inserir_imagem_pagina_inteira(p, png)
        self._quebra_pagina()

    def _capa_nativa(self):
        """Capa em texto nativo do Word (fallback sem Chromium/Playwright)."""
        t = self.tokens
        lau = self.laudo
        for _ in range(3):
            self.doc.add_paragraph()
        self._p("PAULA PACHECO", align=WD_ALIGN_PARAGRAPH.CENTER, size=13,
                bold=True, cor=t["marca"]["azul"], space_after=0)
        self._p("ENGENHARIA E PERÍCIAS", align=WD_ALIGN_PARAGRAPH.CENTER, size=11,
                cor=t["marca"]["cinza"], space_after=24)

        titulo = self.doc.add_paragraph(style="Title")
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo.add_run(f"Laudo de Engenharia\nNº {lau.numero}")

        self._p(TIPO_ROTULO.get(lau.tipo, lau.tipo).upper(),
                align=WD_ALIGN_PARAGRAPH.CENTER, size=11, cor=t["marca"]["terracota"],
                space_after=28)

        for rot, val in (("Cliente", lau.cliente_nome),
                         ("CNPJ/CPF", lau.cliente_cnpj_cpf),
                         ("Endereço", lau.endereco),
                         ("Diligências", " · ".join(lau.datas_diligencia)),
                         ("Acompanhamento", lau.acompanhamento),
                         ("Data de emissão", lau.data_emissao)):
            if val:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"{rot}: ")
                r.bold = True
                p.add_run(val)
                p.paragraph_format.space_after = Pt(2)

        for _ in range(2):
            self.doc.add_paragraph()
        for linha in textos.bloco_credenciais():
            self._p(linha, align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
                    cor=t["marca"]["cinza"], space_after=0)
        self._quebra_pagina()

    # -------------------------------------------------------------- sumário
    def _sumario(self):
        self.doc.add_heading("Sumário", level=1)
        inserir_sumario(self.doc.add_paragraph())
        self._quebra_pagina()

    # --------------------------------------------------- capítulos fixos db
    def _capitulo_fixo(self, chave: str, numero: int):
        titulo, corpo = textos.CAPITULOS_FIXOS[chave]
        self.doc.add_heading(f"{numero}. {titulo}", level=1)
        for par in corpo.split("\n\n"):
            self._p(par)

    # ------------------------------------------------------------- vistoria
    def _vistoria(self, numero: int):
        lau = self.laudo
        self.doc.add_heading(f"{numero}. Vistoria", level=1)
        self._p(
            "A vistoria foi realizada in loco, com registro fotográfico das "
            "anomalias constatadas nos ambientes a seguir relacionados, na(s) "
            f"data(s) de {', '.join(lau.datas_diligencia) or 'diligência'}, "
            f"com acompanhamento de {lau.acompanhamento or 'representante indicado'}."
        )
        if lau.ambientes:
            self._p("Ambientes vistoriados:", bold=True, space_after=2)
            for amb in sorted(lau.ambientes, key=lambda a: a.ordem):
                p = self.doc.add_paragraph(style="List Bullet")
                pav = f" — {amb.pavimento}" if amb.pavimento else ""
                p.add_run(f"{amb.nome}{pav}")
        # fotos sem vínculo com anomalia = registro complementar da vistoria
        avulsas = [f for f in sorted(lau.fotos, key=lambda f: f.ordem) if not f.anomalia_id]
        if avulsas:
            self._p("Registro fotográfico complementar (condições gerais observadas):",
                    bold=True, space_after=2)
            for foto in avulsas:
                self._figura(foto, lau.ambiente(foto.ambiente_id))

    def _vistoria_complemento(self, numero: int):
        """Continuação nativa do cap. Vistoria (após a página-template):
        ambientes vistoriados e registro fotográfico complementar — conteúdo
        de tamanho variável que não cabe na página fixa do design."""
        lau = self.laudo
        avulsas = [f for f in sorted(lau.fotos, key=lambda f: f.ordem)
                   if not f.anomalia_id]
        if not (lau.ambientes or avulsas):
            return False
        self._garantir_miolo()
        self.doc.add_heading(
            f"{numero}.2 — Ambientes vistoriados e registro complementar",
            level=2)
        if lau.ambientes:
            self._p("Ambientes vistoriados:", bold=True, space_after=2)
            for amb in sorted(lau.ambientes, key=lambda a: a.ordem):
                p = self.doc.add_paragraph(style="List Bullet")
                pav = f" — {amb.pavimento}" if amb.pavimento else ""
                p.add_run(f"{amb.nome}{pav}")
        if avulsas:
            self._p("Registro fotográfico complementar (condições gerais "
                    "observadas):", bold=True, space_after=2)
            for foto in avulsas:
                self._figura(foto, lau.ambiente(foto.ambiente_id))
        return True

    # ---------------------------------------------- relatório / diagnóstico
    def _diagnostico(self, numero: int):
        self.doc.add_heading(f"{numero}. Relatório Técnico — Diagnóstico por "
                             "Sistema Construtivo", level=1)
        self._p(
            "Apresenta-se, a seguir, a análise fenomenológica das anomalias "
            "constatadas, organizada por sistema construtivo. Para cada anomalia "
            "registram-se a descrição, o mecanismo de ação, a causa provável, as "
            "consequências, a origem segundo a taxonomia pericial e a "
            "classificação de risco (IBAPE e Matriz GUT)."
        )
        por_sistema: "OrderedDict[str, list]" = OrderedDict()
        for an in sorted(self.laudo.anomalias, key=lambda a: a.ordem):
            por_sistema.setdefault(an.sistema_construtivo, []).append(an)

        for sistema, anomalias in por_sistema.items():
            self.doc.add_heading(self._sentenca(sistema), level=2)
            for an in anomalias:
                self._anomalia(an)

    def _anomalia(self, an):
        amb = self.laudo.ambiente(an.ambiente_id)
        rotulo_amb = f"  ({amb.nome})" if amb else ""
        self.doc.add_heading(f"{an.titulo}{rotulo_amb}", level=3)

        self._campo_estruturado("Descrição fenomenológica", an.descricao_fenomenologica)
        self._campo_estruturado("Mecanismo", an.mecanismo)
        self._campo_estruturado("Causa provável", an.causa_provavel)
        self._campo_estruturado("Consequências", an.consequencias)
        self._campo_estruturado("Origem (taxonomia)", an.origem_taxonomia)

        # Ressalva técnica automática (vício x manutenção indeterminado)
        if risco.deve_gerar_ressalva_vicio(an):
            self._p(textos.RESSALVA_VICIO, italic=True, cor=self.tokens["status"]["ressalva"]["cor"])

        # Mini-tabela de classificação de risco (IBAPE + GUT calculados)
        self._tabela_classificacao(an)

        # Alerta jurídico automático (GR3 / risco a pessoas/estrutura)
        if risco.deve_gerar_alerta_juridico(an):
            p = self._p(textos.ALERTA_JURIDICO, bold=False,
                        cor=self.tokens["grau_risco"]["GR3"]["cor"])
            p.paragraph_format.space_before = Pt(4)

        # Figuras da anomalia (comprimidas + numeradas)
        for foto in self.laudo.fotos_da_anomalia(an.id):
            self._figura(foto, amb)

    def _tabela_classificacao(self, an):
        cab = ["GR (IBAPE)", "G", "U", "T", "G×U×T", "Faixa de Prioridade", "Prazo"]
        tab = self.doc.add_table(rows=2, cols=len(cab))
        tab.alignment = WD_TABLE_ALIGNMENT.CENTER
        tab.style = "Table Grid"
        for j, texto in enumerate(cab):
            cell = tab.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(texto)
            run.bold = True
            run.font.size = Pt(8.5)
            sombrear_celula(cell, hexnum(self.tokens["marca"]["azul"]))
            run.font.color.rgb = RGBColor.from_string("FFFFFF")

        valores = [risco.GR_ROTULO[an.gr], str(an.g), str(an.u), str(an.t),
                   str(an.gut), an.faixa, an.prazo]
        for j, val in enumerate(valores):
            cell = tab.rows[1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sombrear_celula(tab.rows[1].cells[0], risco.GR_SUPERFICIE[an.gr])
        sombrear_celula(tab.rows[1].cells[5], risco.GR_SUPERFICIE[an.gr])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _figura(self, foto, amb):
        # numeração automática pela ORDEM DE APARIÇÃO no documento
        self._fig_seq += 1
        foto.numero_figura = self._fig_seq
        self.dir_img.mkdir(parents=True, exist_ok=True)
        destino = self.dir_img / f"fig{foto.numero_figura:02d}.jpg"
        origem = Path(foto.arquivo) if foto.arquivo else None
        legenda_base = foto.legenda or (amb.nome if amb else "")
        if origem and origem.exists():
            comprimir_imagem(origem, destino)
        else:
            gerar_placeholder(destino, legenda_base or f"Figura {foto.numero_figura}")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(destino), width=Cm(13))

        # legenda usa o ambiente da PRÓPRIA foto; o da anomalia é fallback
        amb = self.laudo.ambiente(foto.ambiente_id) or amb
        ambiente = amb.nome if amb else ""
        descricao = foto.legenda or ""
        texto_leg = f"FIG. {foto.numero_figura:02d}"
        if ambiente:
            texto_leg += f"  {ambiente}:"
        if descricao:
            texto_leg += f" {descricao}"
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(texto_leg)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(hexnum(self.tokens["marca"]["cinza"]))
        cap.paragraph_format.space_after = Pt(10)

    # ------------------------------------------------------ plano/conclusão
    def _plano_intervencao(self, numero: int):
        self.doc.add_heading(f"{numero}. Plano de Intervenção e Recuperação", level=1)
        self._p(
            "As intervenções recomendadas são priorizadas segundo a Matriz GUT, "
            "observando-se os prazos abaixo. Recomenda-se a elaboração de projeto "
            "executivo específico e a execução por responsável técnico habilitado."
        )
        cab = ["Item / Anomalia", "Sistema", "GR", "Prazo Sugerido", "Faixa"]
        anomalias = sorted(self.laudo.anomalias, key=lambda a: -a.gut)
        tab = self.doc.add_table(rows=1 + len(anomalias), cols=len(cab))
        tab.style = "Table Grid"
        self._preencher_cabecalho(tab, cab)
        for i, an in enumerate(anomalias, start=1):
            linha = tab.rows[i].cells
            self._set_cell(linha[0], an.titulo, size=9)
            self._set_cell(linha[1], self._sentenca(an.sistema_construtivo), size=9)
            self._set_cell(linha[2], f"GR{an.gr}", size=9, center=True)
            self._set_cell(linha[3], an.prazo, size=9, center=True)
            self._set_cell(linha[4], an.faixa, size=9, center=True)
            sombrear_celula(linha[2], risco.GR_SUPERFICIE[an.gr])

    def _conclusoes(self, numero: int):
        self.doc.add_heading(f"{numero}. Conclusões", level=1)
        gr3 = [a for a in self.laudo.anomalias if a.gr >= 3]
        self._p(
            f"Foram constatadas {len(self.laudo.anomalias)} anomalias, distribuídas "
            f"entre os sistemas construtivos inspecionados, das quais {len(gr3)} "
            "classificada(s) como de risco alto (GR3), demandando intervenção "
            "prioritária. Recomenda-se a adoção das medidas indicadas no Plano de "
            "Intervenção, respeitados os prazos calculados pela Matriz GUT."
        )

        # Quadro-resumo qualitativo: Item / Grau IBAPE / Prazo Sugerido
        self.doc.add_heading("Quadro-resumo (qualitativo — IBAPE)", level=2)
        cab = ["Item", "Anomalia", "Grau IBAPE", "Prazo Sugerido"]
        anomalias = sorted(self.laudo.anomalias, key=lambda a: a.ordem)
        tab = self.doc.add_table(rows=1 + len(anomalias), cols=len(cab))
        tab.style = "Table Grid"
        self._preencher_cabecalho(tab, cab)
        for i, an in enumerate(anomalias, start=1):
            c = tab.rows[i].cells
            self._set_cell(c[0], str(i), size=9, center=True)
            self._set_cell(c[1], an.titulo, size=9)
            self._set_cell(c[2], risco.GR_ROTULO[an.gr], size=9, center=True)
            self._set_cell(c[3], an.prazo, size=9, center=True)
            sombrear_celula(c[2], risco.GR_SUPERFICIE[an.gr])

        # Matriz GUT quantificada
        self.doc.add_heading("Matriz GUT (quantificada)", level=2)
        cab = ["Item", "G", "U", "T", "G×U×T", "Faixa de Prioridade"]
        tab = self.doc.add_table(rows=1 + len(anomalias), cols=len(cab))
        tab.style = "Table Grid"
        self._preencher_cabecalho(tab, cab)
        for i, an in enumerate(anomalias, start=1):
            c = tab.rows[i].cells
            self._set_cell(c[0], an.titulo, size=9)
            self._set_cell(c[1], str(an.g), size=9, center=True)
            self._set_cell(c[2], str(an.u), size=9, center=True)
            self._set_cell(c[3], str(an.t), size=9, center=True)
            self._set_cell(c[4], str(an.gut), size=9, center=True, bold=True)
            self._set_cell(c[5], an.faixa, size=9, center=True)
            sombrear_celula(c[5], risco.GR_SUPERFICIE[an.gr])

    def _encerramento(self, numero: int):
        self.doc.add_heading(f"{numero}. Encerramento", level=1)
        # contagem aproximada de páginas (o Word recalcula ao abrir via NUMPAGES)
        for par in textos.encerramento(self._paginas_estimadas(), self.laudo.numero):
            self._p(par)
        for _ in range(2):
            self.doc.add_paragraph()
        self._p("_" * 38, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        for linha in textos.bloco_credenciais():
            self._p(linha, align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
                    cor=self.tokens["marca"]["cinza"], space_after=0)

    def _paginas_estimadas(self) -> str:
        return "número indicado no rodapé"

    # ------------------------------------------------------ util de tabela
    def _preencher_cabecalho(self, tab, cabecalho):
        for j, texto in enumerate(cabecalho):
            cell = tab.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(texto)
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
            sombrear_celula(cell, hexnum(self.tokens["marca"]["azul"]))

    def _set_cell(self, cell, texto, size=10, bold=False, center=False):
        cell.text = ""
        run = cell.paragraphs[0].add_run(texto)
        run.font.size = Pt(size)
        run.bold = bold
        if center:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------- montagem
    def gerar(self) -> Path:
        self._estilos()
        configurar_cabecalho_rodape(self.doc, self.tokens, self.laudo.numero)
        self._preparar_paginas()
        tem = self._paginas.__contains__

        self._em_imagem = False
        self._capa()
        self._sumario()
        self._capitulo_fixo("pressupostos", 1)

        if tem("motivacao"):
            self._secao_imagem()
            self._pagina_imagem("motivacao", "2. Motivação")
        else:
            self._garantir_miolo()
            self._capitulo_fixo("motivacao", 2)

        if tem("vistoria"):
            self._secao_imagem()
            self._pagina_imagem("vistoria", "3. Vistoria")
            self._vistoria_complemento(3)
        else:
            self._garantir_miolo()
            self._vistoria(3)

        if tem("metodos"):
            self._secao_imagem()
            self._pagina_imagem("metodos", "4. Métodos e Procedimentos")
        else:
            self._garantir_miolo()
            self._capitulo_fixo("metodos", 4)

        self._garantir_miolo()
        self._diagnostico(5)
        self._plano_intervencao(6)

        if tem("especificacoes"):
            self._secao_imagem()
            self._pagina_imagem("especificacoes",
                                "7. Especificações e Aceite")
        else:
            self._capitulo_fixo("especificacoes", 7)

        self._garantir_miolo()
        self._conclusoes(8)

        if tem("referencial"):
            self._secao_imagem()
            self._pagina_imagem("referencial", "9. Referencial Técnico")
        else:
            self._capitulo_fixo("referencial", 9)

        if tem("encerramento"):
            self._secao_imagem()
            self._pagina_imagem("encerramento", "10. Encerramento")
        else:
            self._garantir_miolo()
            self._encerramento(10)

        atualizar_campos_ao_abrir(self.doc)
        self.saida.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.saida))
        return self.saida


def gerar_laudo(caminho_json: str | Path, saida: str | Path) -> Path:
    laudo = carregar_de_json(caminho_json)
    return GeradorLaudo(laudo, Path(saida)).gerar()


if __name__ == "__main__":
    entrada = sys.argv[1] if len(sys.argv) > 1 else "dados/laudo_50_2026.json"
    destino = sys.argv[2] if len(sys.argv) > 2 else "saida/Laudo-50-2026.docx"
    caminho = gerar_laudo(entrada, destino)
    print(f"Laudo gerado: {caminho}")
